#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成摩罗互娱网络安全整改方案实施进度报告 Word 文档。"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Twips


OUTPUT = "/workspace/doc/reports/摩罗互娱网络安全整改方案实施进度报告_20260729.docx"


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_format(p, space_before=0, space_after=6, line_spacing=1.5, first_line=False, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line:
        pf.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_format(p, space_before=14, space_after=8, line_spacing=1.5)
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=14, bold=True)
    elif level == 2:
        set_paragraph_format(p, space_before=10, space_after=6, line_spacing=1.5)
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=12, bold=True)
    else:
        set_paragraph_format(p, space_before=6, space_after=4, line_spacing=1.5)
        run = p.add_run(text)
        set_run_font(run, name="楷体", size=12, bold=True)
    return p


def add_body(doc, text, first_line=True, space_after=6):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=0, space_after=space_after, line_spacing=1.5, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=12)
    return p


def add_bullet(doc, text, indent=0.74):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=0, space_after=3, line_spacing=1.5, first_line=False)
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run("• " + text)
    set_run_font(run, name="宋体", size=12)
    return p


def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "666666")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def fill_cell(cell, text, bold=False, center=False, size=10.5, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, name="宋体", size=size, bold=bold)
    if fill:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, bold=True, center=True, size=10.5, fill="D9E2F3")
    for r_idx, row in enumerate(rows):
        bg = "F2F2F2" if r_idx % 2 else "FFFFFF"
        for c_idx, val in enumerate(row):
            fill_cell(table.rows[r_idx + 1].cells[c_idx], str(val), center=True, size=10.5, fill=bg)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第 ")
    set_run_font(run, name="宋体", size=9, color=RGBColor(0x66, 0x66, 0x66))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run2 = p.add_run()
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_sep)
    run2._r.append(fld_text)
    run2._r.append(fld_end)
    set_run_font(run2, name="宋体", size=9, color=RGBColor(0x66, 0x66, 0x66))

    run3 = p.add_run(" 页")
    set_run_font(run3, name="宋体", size=9, color=RGBColor(0x66, 0x66, 0x66))


def build():
    doc = Document()

    # 页面设置：A4，适合 3-4 页正式报告
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    # ========== 标题区 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title, space_before=6, space_after=6, line_spacing=1.5)
    run = title.add_run("摩罗互娱网络安全整改方案")
    set_run_font(run, name="黑体", size=20, bold=True)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title2, space_before=0, space_after=10, line_spacing=1.5)
    run = title2.add_run("实施进度报告")
    set_run_font(run, name="黑体", size=20, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(meta, space_before=2, space_after=2, line_spacing=1.3)
    run = meta.add_run("依据：摩罗互娱网络安全整改方案技术服务合同")
    set_run_font(run, name="楷体", size=11)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(meta2, space_before=0, space_after=2, line_spacing=1.3)
    run = meta2.add_run("报告日期：2026年7月29日　　整体进度：约 70%")
    set_run_font(run, name="楷体", size=11, bold=True)

    # ========== 一、概述 ==========
    add_heading_cn(doc, "一、概述", level=1)
    add_body(
        doc,
        "根据《摩罗互娱网络安全整改方案技术服务合同》约定内容，本报告对当前网络安全整改工作的实施进展进行阶段性汇总。"
        "截至本报告日期，整体整改工作进度约达70%。整改工作围绕渗透测试与端口收敛、操作系统升级与核心/大地市服务器迁移、"
        "内网围墙重塑、AI代码审计与高风险接口治理等主线推进，重点降低老互娱生产环境公网暴露面，提升运维体系安全基线。"
        "本报告重点说明已完成事项、关键量化结果、后续迁移与加固计划，并附相关交付材料清单，供项目各方跟踪确认。",
    )

    add_heading_cn(doc, "（一）分项进度一览", level=2)
    add_table(
        doc,
        headers=["整改事项", "当前状态", "完成情况说明"],
        rows=[
            ["渗透测试与端口收敛", "已完成", "2172个端口核查处置完毕，非业务端口清零"],
            ["核心服务器升级迁移", "已完成", "运维/跳板机/监控服务器升级至 Ubuntu 24.04"],
            ["大地市操作系统迁移", "进行中", "7个受攻击大地市中已完成3个"],
            ["内网围墙重塑", "已完成", "非用户访问系统与接口已封锁进内网/VPN"],
            ["AI代码审计与接口治理", "基本完成", "对外PHP已替换为Go；配牌后门仅内网开放待加告警"],
        ],
        col_widths=[4.2, 2.8, 6.7],
    )
    doc.add_paragraph()

    # ========== 二、渗透测试 ==========
    add_heading_cn(doc, "二、渗透测试（已完成）", level=1)
    add_body(
        doc,
        "老互娱生产环境渗透扫描覆盖公网开放端口共计2172个。经逐项核查与处置，非业务端口已全部关闭；"
        "业务端口中未使用端口已完成收敛；仍保留开放的业务端口均为确认为业务所需端口。",
    )

    add_heading_cn(doc, "（一）端口处理情况统计", level=2)
    add_table(
        doc,
        headers=["序号", "统计项", "数量（个）"],
        rows=[
            ["1", "总的非业务端口", "289"],
            ["2", "总的业务端口", "1883"],
            ["3", "非业务端口已关闭", "289"],
            ["4", "业务端口已关闭（未使用）", "187"],
            ["5", "未关闭的非业务端口", "0"],
            ["6", "未关闭的业务端口", "1696"],
        ],
        col_widths=[2.2, 7.5, 4.0],
    )
    doc.add_paragraph()

    add_heading_cn(doc, "（二）服务端口归类关闭统计", level=2)
    add_table(
        doc,
        headers=["序号", "服务类型", "关闭数量（个）"],
        rows=[
            ["1", "php-fpm 服务端口", "116"],
            ["2", "zabbix 服务端口", "88"],
            ["3", "sshd 服务端口", "33"],
            ["4", "clubserver 服务端口", "33"],
        ],
        col_widths=[2.2, 7.5, 4.0],
    )
    doc.add_paragraph()

    add_heading_cn(doc, "（三）关键整改结论", level=2)
    add_bullet(doc, "目前各地市外网已无 PHP 接口开放；公共服对外开放的 PHP 接口均已完成代码审查，存在风险的 PHP 接口已更新为 Go 重写。")
    add_bullet(doc, "Zabbix、clubserver 等无需外网访问的内部服务，均已关闭外网访问。")
    add_bullet(doc, "高风险运维管理服务器的 SSH 端口已全部封入内网，仅允许由内部运维服务器进行管理。")

    add_heading_cn(doc, "（四）相关交付材料", level=2)
    add_body(doc, "未加固前端口开放情况：", first_line=False, space_after=3)
    add_bullet(doc, "地市生产服务器渗透开启端口服务表-复查完成.xlsx")
    add_bullet(doc, "公共服务器公网端口渗透扫结果-2026-06-10.xlsx")
    add_body(doc, "加固后端口情况：", first_line=False, space_after=3)
    add_bullet(doc, "老互娱全量公网端口扫漏对比报告_20260729.xlsx")

    add_heading_cn(doc, "（五）后续建议", level=2)
    add_body(
        doc,
        "渗透测试方案已交付运维团队。建议后续按月执行一次外网端口扫描复查，持续确认公网暴露面无回流、无新增高风险端口开放，"
        "并将扫描结果纳入运维例行安全检查闭环。",
    )

    # ========== 三、操作系统升级 ==========
    add_heading_cn(doc, "三、操作系统升级与服务器迁移", level=1)

    add_heading_cn(doc, "（一）核心服务器升级迁移（已完成）", level=2)
    add_body(
        doc,
        "运维服务器、跳板机服务器、监控服务器均已完成升级并迁移至 Ubuntu 24.04，弃用已失去官方系统漏洞修复支持的 CentOS 7，"
        "从操作系统层面降低已知漏洞长期暴露风险，夯实核心运维基础设施安全基线。",
    )

    add_heading_cn(doc, "（二）大地市操作系统升级与迁移", level=2)
    add_body(
        doc,
        "目前已完成海满、摩罗长春、齐齐哈尔三大地市的操作系统及服务器迁移。铁岭因市场原因靠后处理。"
        "老互娱受攻击地市统计共有7个大地市，迁移进展说明如下。",
    )

    add_body(doc, "今年老互娱受攻击发生地市：新阜新、保定、铁岭、大连、海满、摩罗长春、齐齐哈尔。", first_line=True)
    add_body(doc, "已完成迁移地市：海满、摩罗长春、齐齐哈尔。", first_line=True)

    add_table(
        doc,
        headers=["地市", "迁移状态", "备注"],
        rows=[
            ["海满", "已完成", "操作系统与服务器迁移完成"],
            ["摩罗长春", "已完成", "操作系统与服务器迁移完成"],
            ["齐齐哈尔", "已完成", "操作系统与服务器迁移完成"],
            ["保定", "待迁移", "后续迁移顺序第1位"],
            ["大连", "待迁移", "后续迁移顺序第2位"],
            ["新阜新", "待迁移", "后续迁移顺序第3位"],
            ["铁岭", "待迁移", "因市场原因靠后处理"],
        ],
        col_widths=[3.5, 3.5, 6.7],
    )
    doc.add_paragraph()

    add_body(
        doc,
        "后续迁移顺序建议为：保定 → 大连 → 新阜新 → 铁岭。其余地市可由运维团队规划逐步升级，或与后续降本方案一并统筹讨论推进。",
    )

    # ========== 四、内网围墙 ==========
    add_heading_cn(doc, "四、内网围墙重塑", level=1)
    add_body(
        doc,
        "针对无需面向用户访问的系统与接口，已实施内网化与访问边界收敛，将相关服务封锁进内网及 VPN，降低外部攻击面。"
        "公网端口收敛与对比结果详见《老互娱全量公网端口扫漏对比报告_20260729.xlsx》。",
    )
    add_body(
        doc,
        "通过“业务必需开放、非业务一律关闭、内部服务内网化”的原则，形成更清晰的内外网隔离边界，"
        "使运维管理入口、监控组件、历史遗留内部服务等不再直接暴露于公网。",
    )

    # ========== 五、AI代码审计 ==========
    add_heading_cn(doc, "五、AI代码审计与接口治理", level=1)
    add_bullet(doc, "所有对外 PHP 接口均已替换为 Go 接口，从技术栈层面消除外网 PHP 接口暴露风险。")
    add_bullet(
        doc,
        "Go 与 C++ 接口经审计后，仅存一个用于配牌的后门接口；该接口仅内网开放，并将补充调用告警机制，"
        "一旦出现调用即在群内告警并及时处置。",
    )
    add_body(doc, "详细审计结论见：《Go 与 C++ 扫描报告.docx》。", first_line=True)

    # ========== 六、阶段性总结 ==========
    add_heading_cn(doc, "六、阶段性总结与下一步计划", level=1)
    add_body(
        doc,
        "当前整改工作已完成渗透测试与端口收敛、核心运维基础设施升级迁移、部分受攻击大地市迁移、内网围墙重塑及对外 PHP 接口治理等关键任务，"
        "整体进度约70%。从效果看：公网非业务端口已清零；Zabbix、clubserver 等内部服务已取消外网访问；高风险 SSH 管理入口已全部内网化；"
        "各地市外网已无 PHP 接口开放，相关风险接口已完成 Go 重写替换。上述措施显著压缩了外部攻击面。",
    )
    add_body(
        doc,
        "尚未完成的主要工作集中在受攻击大地市的操作系统升级与服务器迁移。已完成海满、摩罗长春、齐齐哈尔；"
        "保定、大连、新阜新、铁岭仍待按序推进。同时，配牌后门接口的内网调用告警机制需尽快落地，形成发现—告警—处置闭环。",
    )

    add_heading_cn(doc, "（一）下一步重点工作", level=2)
    add_bullet(doc, "按既定顺序推进保定、大连、新阜新、铁岭等地市操作系统升级与服务器迁移。")
    add_bullet(doc, "由运维按月执行外网端口扫描复查，防止高风险端口回流，并将结果纳入例行安全检查。")
    add_bullet(doc, "完善配牌后门接口的内网调用告警与应急处置流程，确保异常调用可及时发现并响应。")
    add_bullet(doc, "结合降本方案，统筹剩余地市升级节奏与资源投入，避免重复建设与资源浪费。")

    add_heading_cn(doc, "（二）附件清单", level=2)
    add_table(
        doc,
        headers=["序号", "附件名称", "用途说明"],
        rows=[
            ["1", "地市生产服务器渗透开启端口服务表-复查完成.xlsx", "加固前地市生产端口基线"],
            ["2", "公共服务器公网端口渗透扫结果-2026-06-10.xlsx", "加固前公共服公网端口基线"],
            ["3", "老互娱全量公网端口扫漏对比报告_20260729.xlsx", "加固前后端口对比与内网围墙结果"],
            ["4", "Go 与 C++ 扫描报告.docx", "AI代码审计与接口治理明细"],
        ],
        col_widths=[1.5, 8.2, 4.0],
    )
    doc.add_paragraph()

    # 结尾
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(end, space_before=14, space_after=4, line_spacing=1.5)
    run = end.add_run("报告编制日期：2026年7月29日")
    set_run_font(run, name="宋体", size=12)

    end2 = doc.add_paragraph()
    end2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(end2, space_before=0, space_after=0, line_spacing=1.5)
    run = end2.add_run("（本报告为阶段性实施进度汇报）")
    set_run_font(run, name="楷体", size=11, color=RGBColor(0x66, 0x66, 0x66))

    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


if __name__ == "__main__":
    build()
